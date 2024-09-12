import openai
from PyPDF2 import PdfReader
from docx import Document
import time
import tiktoken  # Ensure you install this package

# Configure the OpenAI API with your API key
openai.api_key = ""

# Function to estimate token count based on the content using tiktoken library
def count_tokens(text, model="gpt-3.5-turbo-16k"):
    encoding = tiktoken.encoding_for_model(model)
    return len(encoding.encode(text))

# Function to extract text from a PDF
def extract_text_from_pdf(pdf_file_path):
    pdf_reader = PdfReader(pdf_file_path)
    raw_text = ''.join([page.extract_text() for page in pdf_reader.pages])
    return raw_text

# Function to extract text from a DOCX file
def extract_text_from_docx(docx_file_path):
    doc = Document(docx_file_path)
    raw_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    return raw_text

# Function to process the text using OpenAI GPT-3.5 Turbo with token count check
def process_text_with_gpt(text, reference_content, max_tokens=1000):
    prompt = f"""You must refer strictly to the reference template provided below and perform the following tasks:

1. **Contraindications**: Provide **Contraindications** related to that device. If there are no Contraindications, use 3-4 concise sentences related to the product.
2. **Device Description**: Write a detailed description of the device based **solely** on the provided content. 
3. **Warnings and Cautions :Trimming Sentences with 'Risk of**: Extract Each and Every all warnings and cautions.

Reference Template:\n{reference_content}\n
Content to Process:\n{text}
"""

    prompt_token_count = count_tokens(prompt)
    print(f"Prompt token count: {prompt_token_count}")

    # Ensure the total token count stays within the model's limit
    if prompt_token_count + max_tokens > 16385:
        raise ValueError(f"Token count exceeds limit. Reduce input size or request fewer tokens.")

    start_time = time.time()

    # Using the latest API call structure
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo-16k",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.2,
        max_tokens=max_tokens
    )

    response_time = time.time() - start_time
    print(f"Response generation time: {response_time} seconds")

    generated_text = response['choices'][0]['message']['content'] if response['choices'] else None
    response_token_count = count_tokens(generated_text) if generated_text else None
    return generated_text, prompt_token_count, response_token_count, response_time

# Function to save generated response and token counts to a Word document
def save_to_word(generated_text, prompt_token_count, response_token_count, response_time, output_docx_path):
    doc = Document()
    doc.add_paragraph(generated_text)
    doc.add_paragraph(f"\n\nInput Token Count: {prompt_token_count}")
    doc.add_paragraph(f"Output Token Count: {response_token_count}")
    doc.add_paragraph(f"Total Token Count: {prompt_token_count + response_token_count}")
    doc.add_paragraph(f"Response Time: {response_time:.2f} seconds")
    doc.save(output_docx_path)

# Main function to run the entire process with token handling
def extract_and_generate_response(pdf_file_path, reference_docx_path, output_docx_path):
    raw_text = extract_text_from_pdf(pdf_file_path)
    reference_content = extract_text_from_docx(reference_docx_path)

    if raw_text and reference_content:
        try:
            generated_response, prompt_token_count, response_token_count, response_time = process_text_with_gpt(raw_text, reference_content, max_tokens=1000)

            if generated_response:
                save_to_word(generated_response, prompt_token_count, response_token_count, response_time, output_docx_path)
                print(f"Generated response saved to {output_docx_path}")
            else:
                print("Failed to generate a response.")
        except ValueError as ve:
            print(ve)
    else:
        print("No text extracted from one or both documents.")

# Example usage:
pdf_file_path = "Data Input Document 1_User Manual.pdf"
reference_docx_path = "Refernce-Template-Extract.docx"
output_docx_path = "Generated_Response(5).docx"

extract_and_generate_response(pdf_file_path, reference_docx_path, output_docx_path)

