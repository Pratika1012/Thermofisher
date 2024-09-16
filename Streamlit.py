import streamlit as st
from PIL import Image

from docx import Document
import pandas as pd
import pdfplumber
import docx
import os
import io
import base64
import zipfile
import re
import google.generativeai as genai
import fitz 

from docx.shared import Inches


# Title of the app

st.title("   Welcome to Document Generator   ")




genai.configure(api_key="AIzaSyBNKJ5UoqldD8BwNVCwbDHs3GquaZ9OIjM")
generation_config = {
    "temperature": 0.4,
    "top_p": 0.9,
    "top_k": 50,
    "max_output_tokens": 2048,
}

safety_settings = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"}
]
# Initialize the Gemini model
model = genai.GenerativeModel(
    model_name="gemini-pro",
    generation_config=generation_config,
    safety_settings=safety_settings
)







def extract_and_format_tables(file_path):

    # Determine file type based on the file extension
    file_extension = os.path.splitext(file_path)[1].lower()
    
    # Initialize list to store tables as DataFrames
    all_tables = []

    # If the file is a PDF
    if file_extension == '.pdf':
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    # Reconstruct and clean each table
                    max_columns = max(len(row) for row in table)
                    formatted_table = []
                    for row in table:
                        formatted_row = []
                        for cell in row:
                            # Handle empty or merged cells
                            if cell is None or cell.strip() == '':
                                formatted_row.append(None)
                            else:
                                formatted_row.append(cell.strip())
                        # Ensure each row has the same number of columns
                        while len(formatted_row) < max_columns:
                            formatted_row.append('')
                        formatted_table.append(formatted_row)
                    
                    # Convert to DataFrame
                    df = pd.DataFrame(formatted_table[1:], columns=formatted_table[0])
                    all_tables.append(df)
    
    # If the file is a Word document (.docx)
    elif file_extension == '.docx':
        doc = docx.Document(file_path)
        
        for table in doc.tables:
            formatted_table = []
            max_columns = max(len(row.cells) for row in table.rows)
            
            for row in table.rows:
                formatted_row = []
                for cell in row.cells:
                    # Handle empty or merged cells
                    cell_text = cell.text.strip()
                    if not cell_text:
                        formatted_row.append(None)
                    else:
                        formatted_row.append(cell_text)
                
                # Ensure each row has the same number of columns
                while len(formatted_row) < max_columns:
                    formatted_row.append(None)
                
                formatted_table.append(formatted_row)
            
            # Convert to DataFrame
            df = pd.DataFrame(formatted_table[1:], columns=formatted_table[0])  # Use first row as column headers
            all_tables.append(df)
    
    # If the file type is neither PDF nor DOCX
    else:
        raise ValueError("Unsupported file format. Only '.pdf' and '.docx' files are supported.")
    
    return all_tables

def filter_tables_with_column(dataframes, substring):
    filtered_dfs = []
    for df in dataframes:
        if any(substring in col for col in df.columns):
            filtered_dfs.append(df)
    return filtered_dfs

def save_tables_to_excel(tables):
    output = io.BytesIO()
    # Step 2: Save each table as a sheet in an Excel file
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        for i, table in enumerate(tables):
            table.to_excel(writer, sheet_name=f'Table_{i+1}', index=False)

    output.seek(0)
    return output.getvalue()

def generate_excel_download_link(excel_data,name):
    # Encode the Excel file data to base64
    b64 = base64.b64encode(excel_data).decode()
    
    # Create the download link
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.spreadsheetml.sheet;base64,{b64}" download="{name}.xlsx">Click here to download your Excel file</a>'
    return href




def clean_text(text):
    """
    Remove control characters and invalid XML characters from the text.
    """
    cleaned_text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
    return cleaned_text

def extract_pdf_text(pdf_file,page_number=[]):
    """
    Extract text from the PDF while attempting to include all content, including tables.
    """
    # Read the content of the uploaded file into bytes
    with open(pdf_file, "rb") as pdf_file:
        pdf_bytes = pdf_file.read()
    
    # Open the PDF from bytes
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    extracted_text = []
    if len(page_number)==0:
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            page_height = page.rect.height
            blocks = page.get_text("blocks")  # Extract text blocks with positions
            cleaned_lines = []

            for block in blocks:
                block_text = block[4]  # Text content of the block
                block_y0 = block[1]    # Y-coordinate (top) of the block
                block_y1 = block[3]    # Y-coordinate (bottom) of the block

                # Adjust the footer skip logic
                # Including a more flexible range to cover various types of content
                if block_y1 > page_height * 0.9:
                    continue

                # Clean text and add to the lines
                cleaned_lines.append(clean_text(block_text))

            # Combine cleaned lines for the page
            page_text = "\n".join(cleaned_lines)
            extracted_text.append(page_text)

        doc.close()
        return "\n".join(extracted_text)
    else:
        for page_num in page_number:
            page = doc.load_page(page_num)
            page_height = page.rect.height
            blocks = page.get_text("blocks")  # Extract text blocks with positions
            cleaned_lines = []

            for block in blocks:
                block_text = block[4]  # Text content of the block
                block_y0 = block[1]    # Y-coordinate (top) of the block
                block_y1 = block[3]    # Y-coordinate (bottom) of the block

                # Adjust the footer skip logic
                # Including a more flexible range to cover various types of content
                if block_y1 > page_height * 0.9:
                    continue

                # Clean text and add to the lines
                cleaned_lines.append(clean_text(block_text))

            # Combine cleaned lines for the page
            page_text = "\n".join(cleaned_lines)
            extracted_text.append(page_text)

        doc.close()
        return "\n".join(extracted_text)

def extract_text_from_word(docx_file):
    doc = Document(docx_file)
    raw_text = '\n'.join([para.text for para in doc.paragraphs])
    return raw_text
                    
def process_text_with_gemini(text, reference_content):
    prompt = f"""You must refer strictly to the reference template provided below and perform the following tasks:

1. **Contraindications**:
    Provide **Contraindications** of related that Product/Device only...
    If **Contraindications** are not found in the content, then write  "LLM generated response check manually" :
    
2. **Device Description**:
    If you genearting description of own write an manaual
    - Write a detailed description of the device based **solely** on the provided content.
    - The description must include relevant facts, numbers, figures, features, and any other significant information found in the content.
    - Include any potential adverse effects, as well as important warnings and precautions related to the device.
    - The description should be written in **3-4 paragraphs**.

3. **Warnings and Cautions**:
    - **Extract every single warning and caution** from the content and list them separately under the following format:
        - **Warnings List**:
            - WARNING: [extracted warning]
        - **Cautions List**:
            - CAUTION: [extracted caution]
    - **Do not include** the phrase "Risk of ..." in the extracted warnings and cautions.
    - Ensure **every** warning and caution present in the input PDF content is included.

4. **Text Formatting**:
    - After extracting warnings and cautions, ensure that any "Risk of ..." phrases are removed.
    - Ensure the remaining text is grammatically correct and properly format
    
    **Important**: All warnings and cautions should be processed according to these instructions. Ensure the output follows the format and style of the reference template provided.

Reference Template:\n{reference_content}\n
Content to Process:\n{text}
"""

    # Send the prompt to the Gemini model
    response = model.generate_content(prompt)
    
    return response.text if response else None




image_name=[]
def extract_image_page_number(image_labels,reference_content):
    prompt = f""" this is image name list {image_labels} and this is output template:{reference_content}
             your task is to find the image name from list whose mention in output template in image section and returm only and only figure name

"""
        
    # Generate the rewritten text using the correct method
    return model.generate_content([prompt]).text

def extract_image_titles_from_page(page,page_number):
    titles = []
    text = page.get_text("text")
    for line in text.split('\n'):
        match = re.match(r"Figure \d+\. .+", line)
        if match:
            titles.append(f"{match.group(0)}_{page_number}")
    return titles

def extract_images_from_pdf(pdf_path, output_folder):
    # Open the PDF file
    pdf_document = fitz.open(pdf_path)
    image_count = 0

    # Create the output directory if it does not exist
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for page_number in range(len(pdf_document)):
        page = pdf_document.load_page(page_number)
        image_list = page.get_images(full=True)
        titles = extract_image_titles_from_page(page,page_number)
        
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image = Image.open(io.BytesIO(image_bytes))
            
            # Convert CMYK to RGB if needed
            if image.mode == "CMYK":
                image = image.convert("RGB")
            
            # Use the title or a default name if title is not available
            title = f"{titles[img_index]}" if img_index < len(titles) else f"Figure_{page_number + 1}_{img_index + 1}"
            image_name.append(title)
            title = title.replace(':', '')  # Remove invalid characters for filenames
            image_filename = os.path.join(output_folder, f"{title}.png")
            image.save(image_filename)
            image_count += 1

    return image_name,titles

def final_image_output_gemini(text,reference_content):
    prompt = f""" this is contex {text} and this is output template:{reference_content}
             your task is to find the details about mantion topic in template 
        Instructions:\n"
                "1. **Context Specific**: Start by extracting all relevant text  "
                "that mentions 'Control Panel'. Continue with the next pages until no further relevant content is found.\n\n"
                "2. **Maintain Order**: Keep the extracted content in the same order as it appears in the document.\n\n"
                "3. **Include All Descriptions**: If there are any titles, bullet points, or paragraphs "
                "directly related to 'Control Panel', include them fully without omitting any details.\n\n"
                "Extracted Content:\n"
                "4. Output should be in word document format"

"""
        
    # Generate the rewritten text using the correct method
    return model.generate_content([prompt]).text  # Use the correct method name here




# Function to generate a Word document
def save_text_in_document(text,image=None):
    print(text)
    output = io.BytesIO()
    doc = Document()
    
    if image!= None:
        doc.add_heading(image, level=2)
        doc.add_picture(f"ExtractedImages2\{image}.png",width=Inches(6.0))
    doc.add_paragraph(text)
    doc.save(output)
    output.seek(0)  # Reset the buffer position to the beginning
    return output

def generate_word_download_link(doc_data, filename):
    b64 = base64.b64encode(doc_data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}.docx">Click here to download the Word document</a>'
    return href








# Initialize session state variables if they don't exist
if 'selected_feature' not in st.session_state:
    st.session_state.selected_feature = None


    
col1, col2, col3 = st.columns(3)

st.sidebar.header("selected_feature")
with st.sidebar:
    if st.button("Data Extraction"):
        st.session_state.selected_feature = 'generator'

    if st.button("Doc Summarization"):
        st.session_state.selected_feature = 'summarization'

    if st.button("Web Scraping"):
        st.session_state.selected_feature = 'scraping'


# Logic to display content based on button selection

if st.session_state.selected_feature == 'generator':
    st.subheader("📄 Data Extraction")
    st.write("")
    col1, col2, col3 = st.columns(3)
    


    if 'option' not in st.session_state:
        st.session_state.option = "text"

    with col1:
        if st.button("Text Extraction"):
            st.session_state.option = 'text'
    with col2:
        if st.button("Image Extraction"):
            st.session_state.option = 'image'
    with col3:
        if st.button("Table Extraction"):
            st.session_state.option = 'table'

    
    
    if st.session_state.option == 'text': 
        st.write("") 
        st.markdown("**Text Extraction**")
        input_file_text = st.file_uploader(
            "🔽 Upload Document for Text Extraction ",
            accept_multiple_files=True,
            type=['pdf', 'docx']
        )
        
        if input_file_text:
            st.success(f"{len(input_file_text)} file(s) uploaded successfully!")
            
        st.write("")
        st.write("")
        

        # Now show the Output File Uploader
        output_file_text = st.file_uploader(
            "🔽 Upload Output File Template for Text Extraction ",
            type=['docx', 'pdf']
        )
        
        st.write("")
        st.write("")
        
        
        
        if st.button("Submit"):
            for document in input_file_text:
                

                extract_text=extract_pdf_text(document)
                reference_content = extract_text_from_word(output_file_text)

                if extract_text and reference_content:
                     # Process the extracted text with Gemini for generating response
                    generated_response = process_text_with_gemini(extract_text, reference_content)
                
                    if generated_response:
                       text=save_text_in_document(generated_response)
                       
                    else:
                        st.error("Failed to generate a response using Gemini.")
                else:
                    st.error("No text extracted from one or both documents.")


                

                file_name = os.path.splitext(document.name)[0]
                st.markdown(generate_word_download_link(text.getvalue(),file_name), unsafe_allow_html=True)

                # st.markdown(generate_excel_download_link(tables,file_name), unsafe_allow_html=True)
                



    if st.session_state.option == 'image':
        st.write("")
        st.markdown("**Image Extraction**")
        # Input uploader for multiple files
        input_file_image = st.file_uploader(
            "🔽 Upload Document for image Extraction ",
            accept_multiple_files=True,
            type=['pdf', 'docx']
        )
        
        if input_file_image:
            st.success(f"{len(input_file_image)} file(s) uploaded successfully!")
            
        st.write("")
        st.write("")
        

        # Now show the Output File Uploader
        output_file_image = st.file_uploader(
            "🔽 Upload Output File Template for image Extraction ",
            type=['docx', 'pdf']
        )
        
        st.write("")
        st.write("")
        
        
        
        if st.button("Submit"):
            for document in input_file_image:
                #extract table
                

                
                reference_content = extract_text_from_word(output_file_image)

                image_labels,title =extract_images_from_pdf(document.name,"ExtractedImages2")

                generated_response=extract_image_page_number(image_labels,reference_content)
                image = generated_response.replace("-", "")
                image=image.strip()
                
                number = int(image.split('_')[-1])
                
                page_num=[number-1,number,number+1,number+2]
                
                extract_text=extract_pdf_text(document.name,page_num)
                final_text=final_image_output_gemini(extract_text,reference_content)
                
                text=save_text_in_document(final_text,image)

                file_name = os.path.splitext(document.name)[0]

                

                
                st.markdown(generate_word_download_link(text.getvalue(),file_name), unsafe_allow_html=True)


    if st.session_state.option == 'table':
        st.write("")
        st.markdown("**Table Extraction**") 
        # Input uploader for multiple files
        input_file_table = st.file_uploader(
            "🔽 Upload Document for Table Extraction ",
            accept_multiple_files=True,
            type=['pdf', 'docx']
        )
        
        if input_file_table:
            st.success(f"{len(input_file_table)} file(s) uploaded successfully!")
            
        st.write("")
        st.write("")
        

        # Now show the Output File Uploader
        output_file_table = st.file_uploader(
            "🔽 Upload Output File Template for Data Extraction ",
            type=['docx', 'pdf']
        )
        
        st.write("")
        st.write("")
        
        
        
        if st.button("Submit"):
            for document in input_file_table:
                #extract table
                substring = "Technical Data Sheet"
                extract_tables = extract_and_format_tables(document.name)
                

                tables=save_tables_to_excel(extract_tables)

                #extract text
                

                file_name = os.path.splitext(document.name)[0]

                st.markdown(generate_excel_download_link(tables,file_name), unsafe_allow_html=True)
                




if st.session_state.selected_feature == 'summarization':
    st.subheader("📝 Document Summarization")
    input_file2 = st.file_uploader("🔽 Upload Document for Summarization",accept_multiple_files=True, type=['pdf', 'docx', 'txt'])
    
    if input_file2:
        st.success(f"{len(input_file2)} file(s) uploaded successfully!")

    st.write("")
    st.write("")
    st.write("")

    output_file2 = st.file_uploader(
        "🔽 Upload Output File Template for Summarization",
        type=['docx', 'pdf']
    )
    
    if output_file2:
        st.success(f"Output file '{output_file2.name}' uploaded successfully!")




if st.session_state.selected_feature == 'scraping':
    st.subheader("🌐 Document Scraping")
    url_input = st.text_input("🔽 Enter URL for Scraping")
    
    if url_input:
        st.write(f"🕵️‍♂️ Scraping content from: {url_input}")
        # Implement scraping logic here
    
    st.write("")
    st.write("")
    st.write("")

    output_file3 = st.file_uploader(
        "🔽 Upload Output File Template for Scraping",
        type=['docx', 'pdf']
    )
    
    if output_file3:
        st.success(f"Output file '{output_file3.name}' uploaded successfully!")
