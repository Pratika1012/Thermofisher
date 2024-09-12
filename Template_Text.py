# import streamlit as st
# from PyPDF2 import PdfReader
# from docx import Document
# import io
# import os
# import zipfile

# # Initialize session state to track which section is active and store files
# if 'active_section' not in st.session_state:
#     st.session_state.active_section = None  # No section is active initially
# if 'processed_files' not in st.session_state:
#     st.session_state.processed_files = []  # Store processed files for download

# # Functions to handle button clicks and set session state
# def document_generator():
#     st.session_state.active_section = "doc_gen"

# def document_summarization():
#     st.session_state.active_section = "doc_sum"

# def document_scraping():
#     st.session_state.active_section = "doc_scrap"

# # Text extraction function from PDF
# def extract_text_with_pyPDF(pdf_file):
#     pdf_reader = PdfReader(pdf_file)
#     raw_text = ''
#     for page in pdf_reader.pages:
#         text = page.extract_text()
#         if text:
#             raw_text += text
#     return raw_text

# # Text extraction function from Word document
# def extract_text_from_word(docx_file):
#     doc = Document(docx_file)
#     raw_text = '\n'.join([para.text for para in doc.paragraphs])
#     return raw_text

# # Function to generate a Word document
# def generate_word_document(text, output_path):
#     doc = Document()
#     doc.add_paragraph(text)
#     doc.save(output_path)

# # Function to create a zip of all files in a folder
# def zip_folder(folder_path, zip_filename):
#     zip_buffer = io.BytesIO()
#     with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
#         for root, _, files in os.walk(folder_path):
#             for file in files:
#                 file_path = os.path.join(root, file)
#                 zf.write(file_path, arcname=file)
#     zip_buffer.seek(0)
#     return zip_buffer

# # Streamlit app title
# st.title("Document Processing App")

# # Buttons for the three functionalities
# col1, col2, col3 = st.columns(3)
# with col1:
#     st.button("Document Generator", on_click=document_generator)
# with col2:
#     st.button("Document Summarization", on_click=document_summarization)
# with col3:
#     st.button("Document Scraping", on_click=document_scraping)

# # Document Generator Section
# if st.session_state.active_section == "doc_gen":
#     st.subheader("📄 Document Generator")

#     # Input uploader for multiple files (PDF or Word)
#     input_files = st.file_uploader(
#         "🔽 Upload Input Files (PDF or Word)",
#         accept_multiple_files=True,
#         type=['pdf', 'docx']
#     )
    
#     if input_files:
#         # Create a folder to store all generated files (if it doesn't exist)
#         output_folder = "Extract_Processed_files"
#         if not os.path.exists(output_folder):
#             os.makedirs(output_folder)

#         def process_files_and_zip():
#             for uploaded_file in input_files:
#                 # Extract text based on file type (PDF or Word)
#                 if uploaded_file.name.endswith('.pdf'):
#                     extracted_text = extract_text_with_pyPDF(uploaded_file)
#                     output_name = uploaded_file.name.replace('.pdf', '.docx')
#                 elif uploaded_file.name.endswith('.docx'):
#                     extracted_text = extract_text_from_word(uploaded_file)
#                     output_name = uploaded_file.name  # Keep the same name for Word files
                
#                 # Define output path
#                 output_path = os.path.join(output_folder, output_name)

#                 # Generate output Word file with the extracted text and store it
#                 generate_word_document(extracted_text, output_path)

#             # Once files are processed, zip the folder
#             zip_buffer = zip_folder(output_folder, 'processed_files.zip')

#             # Provide the zip file for download
#             st.download_button(
                
#                 label="Download All Processed Documents",
#                 data=zip_buffer,
                
#                 file_name="Extarct_Processed_files.zip",
#                 mime="application/zip"
#             )

#         # Show a button that processes the files and generates the zip
#         if st.button("Process and Download Documents"):
#             process_files_and_zip()
           
            
        
#         # st.success(f"All files processed and stored in '{output_folder}' folder.")

#         # Upload optional output file
#     output_file = st.file_uploader(
#             "🔽 Upload Output File (Optional)",
#             type=['docx', 'pdf']
#         )
#     if output_file:
#             st.success(f"Output file '{output_file.name}' uploaded successfully!")

#     # Button to download all processed files as a zip
    
         

# # Document Summarization Section (for future use)
# if st.session_state.active_section == "doc_sum":
#     st.subheader("📝 Document Summarization")
#     uploaded_file = st.file_uploader("🔽 Upload Document for Summarization", accept_multiple_files=True,
#                                      type=['pdf', 'docx', 'txt'])
#     if uploaded_file:
#         st.success(f"Files '{len(uploaded_file)}' uploaded successfully!")
    
#     else:
#         st.info("Please upload one or more input files to proceed.")
#         # Implement summarization logic here

# # Document Scraping Section (for future use)
# if st.session_state.active_section == "doc_scrap":
#     st.subheader("🌐 Document Scraping")
#     url_input = st.text_input("🔽 Enter URL for Scraping")
    
#     if url_input:
#         st.write(f"🕵️‍♂️ Scraping content from: {url_input}")
#         # Implement scraping logic here



import streamlit as st
from docx import Document
import google.generativeai as genai
import io
import os
import zipfile
import re
import fitz  # PyMuPDF for PDF handling

genai.configure(api_key="AIzaSyBNKJ5UoqldD8BwNVCwbDHs3GquaZ9OIjM")

# Set up the model configuration
generation_config = {
    "temperature": 0.3,
    "top_p": 0.9,
    "top_k": 50,
    "max_output_tokens": 2048,
}

safety_settings = [
    {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"}
]

# Initialize the Gemini model
model = genai.GenerativeModel(
    model_name="gemini-pro",
    generation_config=generation_config,
    safety_settings=safety_settings
)

# Initialize session state to track which section is active and store files
if 'active_section' not in st.session_state:
    st.session_state.active_section = None  # No section is active initially
if 'processed_files' not in st.session_state:
    st.session_state.processed_files = []  # Store processed files for download

# Functions to handle button clicks and set session state
def document_generator():
    st.session_state.active_section = "doc_gen"

def document_summarization():
    st.session_state.active_section = "doc_sum"

def document_scraping():
    st.session_state.active_section = "doc_scrap"

# Text cleaning function to remove control characters
def clean_text(text):
    """
    Remove control characters and invalid XML characters from the text.
    """
    cleaned_text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
    return cleaned_text

# Text extraction function from PDF using PyMuPDF (fitz)
def extract_pdf_text(pdf_file):
    """
    Extract text from the PDF while attempting to include all content, including tables.
    """
    # Read the content of the uploaded file into bytes
    pdf_bytes = pdf_file.read()
    
    # Open the PDF from bytes
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    extracted_text = []

    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        page_height = page.rect.height
        blocks = page.get_text("blocks")  # Extract text blocks with positions
        cleaned_lines = []

        for block in blocks:
            block_text = block[4]  # Text content of the block
            block_y0 = block[1]    # Y-coordinate (top) of the block
            block_y1 = block[3]    # Y-coordinate (bottom) of the block

            # Adjust the footer skip logic
            # Including a more flexible range to cover various types of content
            if block_y1 > page_height * 0.9:
                continue

            # Clean text and add to the lines
            cleaned_lines.append(clean_text(block_text))

        # Combine cleaned lines for the page
        page_text = "\n".join(cleaned_lines)
        extracted_text.append(page_text)

    doc.close()
    return "\n".join(extracted_text)


# Text extraction function from Word document
def extract_text_from_word(docx_file):
    doc = Document(docx_file)
    raw_text = '\n'.join([para.text for para in doc.paragraphs])
    return raw_text

# Function to generate a Word document
def generate_word_document(text, output_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_path)



def process_text_with_gemini(text, reference_content):
    prompt = f"""You must refer strictly to the reference template provided below and perform the following tasks:

1. **Contraindications**:
    Provide **Contraindications** of related that Product/Device only...
    If **Contraindications** are not found in the content,generate LLM response related to that device.
    
2. **Device Description**:
     If **Device Description** are not found in the content,generate LLM response related to that device.
    - Write a detailed description of the device based **solely** on the provided content.
    - The description must include relevant facts, numbers, figures, features, and any other significant information found in the content.
    - Include any potential adverse effects, as well as important warnings and precautions related to the device.
    - The description should be written in **3-4 paragraphs**.

3. **Warnings and Cautions**:
    - **Extract every single warning and caution** from the content and list them separately under the following format:
        - **Warnings List**:
            - WARNING: [extracted warning]
        - **Cautions List**:
            - CAUTION: [extracted caution]
    - **Do not include** the phrase "Risk of ..." in the extracted warnings and cautions.
    - Ensure **every** warning and caution present in the input PDF content is included.

4. **Text Formatting**:
    - After extracting warnings and cautions, ensure that any "Risk of ..." phrases are removed.
    - Ensure the remaining text is grammatically correct and properly format
    **Important**: All warnings and cautions should be processed according to these instructions. Ensure the output follows the format and style of the reference template provided.

Reference Template:\n{reference_content}\n
Content to Process:\n{text}
"""

    # Send the prompt to the Gemini model
    response = model.generate_content(prompt)
    
    return response.text if response else None

# Function to save generated response to a Word document
def save_to_word(generated_text, output_docx_path):
    doc = Document()
    doc.add_paragraph(generated_text)
    doc.save(output_docx_path)

# Main function to run the entire process
def extract_and_generate_response(pdf_file, reference_file):
    # Extract text from the main PDF
    raw_text = extract_pdf_text(pdf_file)
    
    # Extract text from the reference DOCX file
    reference_content = extract_text_from_word(reference_file)
    
    if raw_text and reference_content:
        # Process the extracted text with Gemini for generating response
        generated_response = process_text_with_gemini(raw_text, reference_content)
        
        if generated_response:
            # Save the generated response to a Word document
            output_path = "Generated_Response.docx"
            save_to_word(generated_response, output_path)
            return output_path
        else:
            st.error("Failed to generate a response using Gemini.")
    else:
        st.error("No text extracted from one or both documents.")

# Function to create a zip of all files in a folder
def zip_folder(folder_path, zip_filename):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for root, _, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                zf.write(file_path, arcname=file)
    zip_buffer.seek(0)
    return zip_buffer

# Streamlit app title
st.title("Document Processing App")

# Buttons for the three functionalities
col1, col2, col3 = st.columns(3)
with col1:
    st.button("Document Generator", on_click=document_generator)
with col2:
    st.button("Document Summarization", on_click=document_summarization)
with col3:
    st.button("Document Scraping", on_click=document_scraping)

# Document Generator Section
if st.session_state.active_section == "doc_gen":
    st.subheader("📄 Document Generator")

    # Input uploader for multiple files (PDF or Word)
    input_files = st.file_uploader(
        "🔽 Upload Input Files (PDF or Word)",
        accept_multiple_files=True,
        type=['pdf', 'docx']
    )
    
    # Reference template uploader
    reference_file = st.file_uploader(
        "🔽 Upload Reference Template (DOCX)",
        type='docx'
    )

    if input_files and reference_file:
        # Create a folder to store all generated files (if it doesn't exist)
        output_folder = "Extract_Processed_files"
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        def process_files_and_zip():
            for uploaded_file in input_files:
                if uploaded_file.name.endswith('.pdf'):
                    extracted_text = extract_pdf_text(uploaded_file)
                    output_name = uploaded_file.name.replace('.pdf', '.docx')
                elif uploaded_file.name.endswith('.docx'):
                    extracted_text = extract_text_from_word(uploaded_file)
                    output_name = uploaded_file.name  # Keep the same name for Word files

                # Define output path
                output_path = os.path.join(output_folder, output_name)

                # Generate output Word file with the extracted text and store it
                generate_word_document(extracted_text, output_path)

            # Zip the folder and provide download
            zip_buffer = zip_folder(output_folder, 'processed_files.zip')
            st.download_button(
                label="Download All Processed Documents",
                data=zip_buffer,
                file_name="Extract_Processed_files.zip",
                mime="application/zip"
            )

        # Show a button that processes the files and generates the zip
        if st.button("Process and Download Documents"):
            process_files_and_zip()

        # Process the extracted text with Gemini for generating response
        if st.button("Generate Response with Reference Template"):
            output_path = extract_and_generate_response(input_files[0], reference_file)
            if output_path:
                st.download_button(
                    label="Download Generated Response",
                    data=open(output_path, 'rb').read(),
                    file_name=output_path,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    # Optional output file uploader
    # output_file = st.file_uploader(
    #     "🔽 Upload Output File (Optional)",
    #     type=['docx', 'pdf']
    # )
    # if output_file:
    #     st.success(f"Output file '{output_file.name}' uploaded successfully!")

# Document Summarization Section (for future use)
if st.session_state.active_section == "doc_sum":
    st.subheader("📝 Document Summarization")
    uploaded_file = st.file_uploader("🔽 Upload Document for Summarization", accept_multiple_files=True,
                                     type=['pdf', 'docx', 'txt'])
    if uploaded_file:
        st.success(f"Files '{len(uploaded_file)}' uploaded successfully!")
    else:
        st.info("Please upload one or more input files to proceed.")
        # Implement summarization logic here

# Document Scraping Section (for future use)
if st.session_state.active_section == "doc_scrap":
    st.subheader("🌐 Document Scraping")
    url_input = st.text_input("🔽 Enter URL for Scraping")
    
    if url_input:
        st.write(f"🕵️‍♂️ Scraping content from: {url_input}")
        # Implement scraping logic here
