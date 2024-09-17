import openai
import tiktoken
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
import time
from langchain import OpenAI, LLMChain
from langchain.prompts import PromptTemplate
import streamlit as st
import re
import fitz

openai.api_key =""

llm = OpenAI(model="gpt-4o", openai_api_key="openai.api_key")


template = PromptTemplate(
    input_variables=["reference_content", "text"],
    template= """You must refer strictly to the reference template provided below and perform the following tasks:

    
1. **Contraindications**:
    Provide **Contraindications** of related that device if there is not Contraindications use the using 3-4 concise sentences associated with the product.
    
2. **Device Description**:
    - Write a detailed description of the device based **solely** on the provided content.
    - The description must include relevant facts, numbers, figures, features, and any other significant information found in the content.
    - Include any potential adverse effects, as well as important warnings and precautions related to the device.
    - The description should be written in **3-4 paragraphs.

3. **Extract Each and Every sentence of WARNING AND CAUTIONS LIST MENTIONED
    - **Before** listing the warnings and cautions, carefully check every sentence in these sections.*. 
    - After removing "Risk of...", the remaining sentence must still be grammatically correct.
    - **Example**:
        - Original: "WARNING: Risk of Shock. Your unit must be properly grounded in conformity with national and local electrical codes."
        - After trimming: "WARNING: Your unit must be properly grounded in conformity with national and local electrical codes."

**Important**: Ensure that all warnings and cautions are processed this way, and the output strictly follows the format, structure, and style of the reference template.

Reference Template:\n{reference_content}\n
Content to Process:\n{text}
"""
)

chain = LLMChain(
    llm=llm,
    prompt=template
)

# Function to estimate token count based on the content using tiktoken library
def count_tokens(text, model="gpt-4o"):
    encoding = tiktoken.encoding_for_model(model)
    return len(encoding.encode(text))



# Define functions for text extraction
def clean_text(text):
    """
    Remove control characters and invalid XML characters from the text.
    """
    cleaned_text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
    return cleaned_text

def extract_pdf_text(pdf_file,page_number=[]):
    """
    Extract text from the PDF while attempting to include all content, including tables.
    """
    # Read the content of the uploaded file into bytes
    with open(pdf_file, "rb") as pdf_file:
        pdf_bytes = pdf_file.read()
    
    # Open the PDF from bytes
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    extracted_text = []
    if len(page_number)==0:
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            page_height = page.rect.height
            blocks = page.get_text("blocks")  # Extract text blocks with positions
            cleaned_lines = []

            for block in blocks:
                block_text = block[4]  # Text content of the block
                block_y0 = block[1]    # Y-coordinate (top) of the block
                block_y1 = block[3]    # Y-coordinate (bottom) of the block

                # Adjust the footer skip logic
                # Including a more flexible range to cover various types of content
                if block_y1 > page_height * 0.9:
                    continue

                # Clean text and add to the lines
                cleaned_lines.append(clean_text(block_text))

            # Combine cleaned lines for the page
            page_text = "\n".join(cleaned_lines)
            extracted_text.append(page_text)

        doc.close()
        return "\n".join(extracted_text)
    else:
        for page_num in page_number:
            page = doc.load_page(page_num)
            page_height = page.rect.height
            blocks = page.get_text("blocks")  # Extract text blocks with positions
            cleaned_lines = []

            for block in blocks:
                block_text = block[4]  # Text content of the block
                block_y0 = block[1]    # Y-coordinate (top) of the block
                block_y1 = block[3]    # Y-coordinate (bottom) of the block

                # Adjust the footer skip logic
                # Including a more flexible range to cover various types of content
                if block_y1 > page_height * 0.9:
                    continue

                # Clean text and add to the lines
                cleaned_lines.append(clean_text(block_text))

            # Combine cleaned lines for the page
            page_text = "\n".join(cleaned_lines)
            extracted_text.append(page_text)

        doc.close()
        return "\n".join(extracted_text)
    
def extract_text_from_word(docx_file):
    """
    Extract text from a DOCX file.
    """
    doc = Document(docx_file)
    raw_text = '\n'.join([para.text for para in doc.paragraphs])
    return clean_text(raw_text)


def process_text_with_langchain(text, reference_content, max_tokens=1000):
    # Create prompt using the template
    prompt = template.format(reference_content=reference_content, text=text)
    
    # Calculate prompt token count
    prompt_token_count = count_tokens(prompt)
    print(f"Prompt token count: {prompt_token_count}")

    # Ensure the total token count stays within the model's limit
    if prompt_token_count + max_tokens > 16385:  # Adjust the limit based on your model
        raise ValueError(f"Token count exceeds limit. Reduce input size or request fewer tokens.")
    
    start_time = time.time()

    # Using LangChain's LLMChain to process the prompt
    result = chain.run(
        reference_content=reference_content,
        text=text,
        max_tokens=max_tokens
    )
    
    # Calculate response token count
    response_token_count = count_tokens(result)
    response_time = time.time() - start_time
    
    return result, prompt_token_count, response_token_count, response_time

# Streamlit UI for file upload
def save_to_word(text):
    """
    Save the text to a DOCX file and return the BytesIO object.
    """
    doc = Document()
    doc.add_paragraph(text)
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Streamlit UI
# Streamlit UI
st.title("Welcome to Document Generator")

pdf_file = st.file_uploader("Upload PDF Document", type=["pdf"])
word_file = st.file_uploader("Upload Word Document (Reference Template)", type=["docx"])

if pdf_file and word_file:
    raw_text = extract_pdf_text(pdf_file.name)
    reference_content = extract_text_from_word(word_file)
    
    if raw_text and reference_content:
        start_time = time.time()
        result, prompt_token_count, response_token_count, response_time = process_text_with_langchain(raw_text, reference_content)

        st.subheader('Generated Response')
        st.text_area("Response:", value=result, height=300)

        st.write(f"Prompt Token Count: {prompt_token_count}")
        st.write(f"Response Token Count: {response_token_count}")
        st.write(f"Response Generation Time: {response_time:.2f} seconds")
        
        output_docx_path = "Generated_Response_Streamlit.docx"
        save_to_word(result).seek(0)
        st.download_button(
            label="Download Generated Response",
            data=save_to_word(result),
            file_name=output_docx_path,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.success(f"Response saved to {output_docx_path}")

    else:
        st.error("No text extracted from one or both documents.")
